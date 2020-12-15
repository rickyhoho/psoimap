from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
import xlwings as xw
import re

excel_file = input("input excel file name:")
question_num = int(input("input number of question:"))

excel_list = []

app = xw.App(visible = True, add_book = False)
wb = xw.Book(excel_file)
sht = wb.sheets[0]

for i in range(2, question_num + 2):
    excel_list.append([])
    for j in range(105, 114):
        excel_list[i - 2].append(sht.range(chr(j) + str(i)).value)
    print(excel_list[i - 2])


doc = Document()    
doc.add_heading('2020hkpsoi_heat_question200',0)
for i in range(question_num):
    question = re.split(r'<(.*?)>', excel_list[i][0])
    for j in range(len(question)):
        if question[j].endswith(".png"):
            if(len(question) > 3):
                doc.add_picture("./2020hkpsoi_heat_question200/" + question[j], width = Inches(1.5))
            else:
                doc.add_picture("./2020hkpsoi_heat_question200/" + question[j], width = Inches(3.7))
        elif j == 0:
            doc.add_paragraph(question[j], style = 'List Number')
        else:
            doc.add_paragraph(question[j])
            
    for j in range(1,5):
        p = doc.add_paragraph()
        run = p.add_run(chr(65 + j - 1) + '.  ')
        if excel_list[i][j * 2] == 10:
            run.font.color.rgb = RGBColor(255,0,0)
        else:
            run.font.color.rgb = RGBColor(0,0,0)
        answer = re.split(r'<(.*?)>', str(excel_list[i][j * 2 - 1]))
        for k in range(len(answer)):
            if answer[k].endswith(".png"):
                doc.add_picture("./2020hkpsoi_heat_question200/" + answer[k], height = Inches(1.6), width = Inches(1.6))
            else:
                run = p.add_run(answer[k])
                if excel_list[i][j * 2] == 10:
                    run.font.color.rgb = RGBColor(255,0,0)
                else:
                    run.font.color.rgb = RGBColor(0,0,0)
    if i != question_num - 1:
        doc.add_page_break()
    print(i)
doc.save('2020hkpsoi_heat_question200.docx')